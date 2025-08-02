"""
Document Processing Utilities Module
===================================

FUNCTION: Core utilities for Word document manipulation and text/image processing.

RESPONSIBILITIES:
- Text replacement in Word documents (handles split text across runs)
- Image insertion and replacement in documents
- Table creation and insertion after specific paragraphs
- File upload handling with secure filename generation
- Document structure manipulation (paragraphs, tables, images)
- Annexure image insertion with caption support

KEY FUNCTIONS:
- find_and_replace_text(): Replaces text placeholders in document
- find_and_replace_image(): Replaces text with images
- insert_table_after(): Creates tables in document structure
- save_uploaded_file(): Securely handles file uploads
- insert_annexure_images(): Adds annexure sections with images

FEATURES:
- Handles text split across multiple runs in Word documents
- Automatic file naming with timestamps
- Secure filename sanitization
- Image resizing and positioning
- Table structure creation with proper XML elements

Document processing utilities for Word document manipulation.
"""
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from werkzeug.utils import secure_filename
import sys
sys.path.append('..')
from config import Config
import time


def find_and_replace_text(doc, old_text, new_text):
    """Finds and replaces text in paragraphs and tables, even if split across runs."""
    def replace_in_paragraph(paragraph):
        full_text = ''.join(run.text for run in paragraph.runs)
        if old_text not in full_text:
            return
        new_full_text = full_text.replace(old_text, new_text)
        for run in paragraph.runs:
            run.text = ''
        if paragraph.runs:
            paragraph.runs[0].text = new_full_text
        else:
            paragraph.add_run(new_full_text)

    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)


def find_and_replace_image(doc, placeholder_text, image_path, width=None):
    """Finds a placeholder in a table cell and replaces it with an image. Returns the cell."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder_text in cell.text:
                    cell.text = ''
                    p = cell.paragraphs[0]
                    run = p.add_run()
                    run.add_picture(image_path, width=width)
                    return cell  # Return the cell for caption insertion
    return None


def insert_paragraph_after(paragraph):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
    return new_para


def insert_table_after(paragraph, rows, cols):
    """Insert a table after the given paragraph."""
    tbl = OxmlElement('w:tbl')
    
    # Table properties
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'auto')
    tblW.set(qn('w:w'), '0')
    tblPr.append(tblW)
    tbl.append(tblPr)
    
    # Create rows and cells
    for i in range(rows):
        tr = OxmlElement('w:tr')
        for j in range(cols):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'auto')
            tcPr.append(tcW)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    
    paragraph._element.addnext(tbl)
    return docx.table.Table(tbl, paragraph._parent)


def save_uploaded_file(file, upload_folder=None):
    """Save uploaded file and return the file path."""
    if upload_folder is None:
        upload_folder = Config.UPLOAD_FOLDER
    
    if file and file.filename:
        filename = secure_filename(file.filename)
        timestamp = str(int(time.time()))
        filename = f"{timestamp}_{filename}"
        file_path = os.path.join(upload_folder, filename)
        file.save(file_path)
        return file_path
    return None


def insert_annexure_images(doc, images, captions, placeholder, image_width=Cm(12), image_height=Cm(20)):
    """Insert annexure images one per page at the placeholder location."""
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, '')
            insert_after = para
            
            for i, img_path in enumerate(images):
                p_img = insert_paragraph_after(insert_after)
                p_img.paragraph_format.alignment = 1  # Center
                run = p_img.add_run()
                run.add_picture(img_path, width=image_width, height=image_height)
                
                # Insert caption if present
                if captions[i]:
                    p_cap = insert_paragraph_after(p_img)
                    p_cap.add_run(captions[i])
                    p_cap.paragraph_format.alignment = 1  # Center
                    p_cap.runs[0].font.size = Pt(10)
                    insert_after = p_cap
                else:
                    insert_after = p_img
                
                # Page break after each image except the last
                if i < len(images) - 1:
                    p_break = insert_paragraph_after(insert_after)
                    p_break.add_run().add_break(docx.text.run.WD_BREAK.PAGE)
                    insert_after = p_break
            
            # Always insert a page break after the last annexure image/caption
            final_break = insert_paragraph_after(insert_after)
            final_break.add_run().add_break(docx.text.run.WD_BREAK.PAGE)
            break  # Only replace the first occurrence


def update_table_of_contents(doc):
    """Update all fields in the document, including table of contents."""
    # Find all field codes in the document
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.xml.find('w:fldChar') != -1:
                # This run contains a field
                run._element.getparent().set(qn('w:dirty'), 'true')
    
    # Find and update TOC fields specifically
    for element in doc.element.body.iter():
        if element.tag.endswith('fldChar'):
            # Mark the field as dirty so it updates when opened
            fld_char_type = element.get(qn('w:fldCharType'))
            if fld_char_type == 'begin':
                # Find the corresponding field instruction
                next_element = element.getnext()
                while next_element is not None:
                    if next_element.tag.endswith('instrText'):
                        if 'TOC' in next_element.text:
                            # Mark TOC field as dirty
                            element.getparent().getparent().set(qn('w:dirty'), 'true')
                        break
                    next_element = next_element.getnext()


def add_update_fields_instruction(doc):
    """Add instruction to update fields when document is opened."""
    # Add document settings to force field updates
    settings = doc.settings
    settings_element = settings._element
    
    # Create updateFields setting
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings_element.append(update_fields)
    
    return doc
