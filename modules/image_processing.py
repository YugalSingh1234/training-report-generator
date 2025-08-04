"""
Image Processing Utilities Module
================================

FUNCTION: Specialized utilities for handling gallery and annexure images in Word documents.

RESPONSIBILITIES:
- Gallery image table creation with multiple images per row
- Annexure image processing and insertion
- Image caption management
- Form data extraction for image-related fields
- Image layout and positioning in document tables
- Page break insertion after image sections

KEY FUNCTIONS:
- insert_gallery_table(): Creates photo gallery tables with images and captions
- get_annexure_images_and_captions(): Extracts annexure data from form
- insert_annexure_images(): Inserts annexure sections with proper formatting

FEATURES:
- Configurable images per row in gallery tables
- Automatic image resizing to specified dimensions
- Caption text placement below images
- Support for multiple annexure sections (Annexure I, II, III, etc.)
- Proper table cell formatting and structure
- Page break management after image sections

Image processing utilities for gallery and annexure images.
"""
from docx.shared import Inches, Pt, Cm
import docx
import docx.oxml.shared
from .document_utils import insert_paragraph_after, insert_table_after, save_uploaded_file


def insert_gallery_table(doc, images, captions, images_per_row=2, image_width=Cm(8.13), placeholder='{{GALLERY_TABLE}}'):
    """
    Inserts tables at the given placeholder with images and captions.
    Each page displays 6 images (2 per row, 3 rows) with equal alignment.
    Images are sized to 8.13cm width × 5.81cm height.
    """
    for para in doc.paragraphs:
        if placeholder in para.text:
            # Remove placeholder text
            para.text = para.text.replace(placeholder, '')
            
            num_images = len(images)
            if num_images == 0:
                return  # No images to insert
            
            images_per_page = 6  # 2 columns × 3 rows = 6 images per page
            rows_per_page = 3    # Fixed 3 rows per page
            insert_after = para
            
            # Process images in batches of 6 (one page at a time)
            for page_start in range(0, num_images, images_per_page):
                page_end = min(page_start + images_per_page, num_images)
                page_images = images[page_start:page_end]
                page_captions = captions[page_start:page_end]
                
                # Calculate actual rows needed for this page
                page_num_images = len(page_images)
                actual_rows = min(rows_per_page, (page_num_images + images_per_row - 1) // images_per_row)
                
                # Insert table for this page
                table = insert_table_after(insert_after, actual_rows, images_per_row)
                
                # Insert images and captions into the table
                img_idx = 0
                for row in table.rows:
                    for cell in row.cells:
                        if img_idx < page_num_images and page_images[img_idx]:
                            # Clear cell and add image
                            cell.text = ''
                            p_img = cell.paragraphs[0]
                            p_img.alignment = 1  # Center alignment for image
                            run = p_img.add_run()
                            run.add_picture(page_images[img_idx], width=image_width, height=Cm(5.81))
                            
                            # Add caption if exists
                            if page_captions[img_idx]:
                                p_caption = cell.add_paragraph(page_captions[img_idx])
                                p_caption.alignment = 1  # Center alignment
                                p_caption.runs[0].font.size = Pt(10)
                                p_caption.runs[0].font.bold = True
                            
                            # Add cell padding and formatting
                            cell.vertical_alignment = 1  # Center vertical alignment
                            
                            img_idx += 1
                
                # Add page break after each page of images (except the last page)
                if page_end < num_images:
                    # Insert page break after the table
                    table_element = table._element
                    parent = table_element.getparent()
                    new_para_element = docx.oxml.shared.OxmlElement('w:p')
                    new_run_element = docx.oxml.shared.OxmlElement('w:r')
                    new_break_element = docx.oxml.shared.OxmlElement('w:br')
                    new_break_element.set(docx.oxml.shared.qn('w:type'), 'page')
                    new_run_element.append(new_break_element)
                    new_para_element.append(new_run_element)
                    parent.insert(parent.index(table_element) + 1, new_para_element)
                    
                    # Update insert_after to the new paragraph for next page
                    insert_after = table
                else:
                    # For the last page, just add the page break without updating insert_after
                    table_element = table._element
                    parent = table_element.getparent()
                    new_para_element = docx.oxml.shared.OxmlElement('w:p')
                    new_run_element = docx.oxml.shared.OxmlElement('w:r')
                    new_break_element = docx.oxml.shared.OxmlElement('w:br')
                    new_break_element.set(docx.oxml.shared.qn('w:type'), 'page')
                    new_run_element.append(new_break_element)
                    new_para_element.append(new_run_element)
                    parent.insert(parent.index(table_element) + 1, new_para_element)
            
            return


def get_annexure_images_and_captions(prefix, request):
    """Get annexure images and captions from form data."""
    images = []
    captions = []
    i = 1
    while True:
        img = save_uploaded_file(request.files.get(f'{prefix}_image_{i}'))
        cap = request.form.get(f'{prefix}_caption_{i}', '')
        if not img:
            break
        images.append(img)
        captions.append(cap)
        i += 1
    return images, captions


def insert_annexure_images(doc, images, captions, placeholder, image_width=Cm(15), image_height=Cm(20), add_final_page_break=True):
    """
    Inserts each image (with caption) at the given placeholder, one per paragraph, sized to fit within page margins.
    Optionally inserts a page break after the last annexure image.
    Images are centered with portrait proportions: 15cm width × 20cm height.
    """
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, '')
            insert_after = para
            for i, img_path in enumerate(images):
                # Insert image with center alignment
                p_img = insert_paragraph_after(insert_after)
                p_img.alignment = 1  # Center alignment
                run = p_img.add_run()
                run.add_picture(img_path, width=image_width, height=image_height)
                
                # Insert caption if present
                if captions[i]:
                    p_cap = insert_paragraph_after(p_img)
                    p_cap.add_run(captions[i])
                    p_cap.paragraph_format.alignment = 1  # Center
                    p_cap.runs[0].font.size = Pt(11)
                    p_cap.runs[0].font.bold = True
                    # Add some spacing after caption
                    p_cap.paragraph_format.space_after = Pt(12)
                    insert_after = p_cap
                else:
                    insert_after = p_img
                # Page break after each image except the last
                if i < len(images) - 1:
                    p_break = insert_paragraph_after(insert_after)
                    p_break.add_run().add_break(docx.text.run.WD_BREAK.PAGE)
                    insert_after = p_break
            
            # Add page break after the last annexure image only if specified
            if add_final_page_break:
                final_break = insert_paragraph_after(insert_after)
                final_break.add_run().add_break(docx.text.run.WD_BREAK.PAGE)
            
            break  # Only replace the first occurrence
