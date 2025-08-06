"""
Chart Processing Utilities Module
=================================

FUNCTION: Specialized utilities for generating feedback charts using matplotlib.

RESPONSIBILITIES:
- Generate horizontal bar charts for feedback questions
- Process feedback data from form submissions
- Create charts with proper styling and colors
- Save charts as images for document insertion
- Handle multiple feedback questions
- Manage chart layout and formatting

KEY FUNCTIONS:
- generate_feedback_charts(): Creates bar charts from feedback data
- process_feedback_data(): Extracts and validates feedback form data
- create_chart_image(): Generates individual chart images
- insert_charts_in_document(): Inserts generated charts into Word document

FEATURES:
- Professional chart styling with custom colors
- Horizontal bar charts for better readability
- Automatic chart sizing and scaling
- Chart export in high resolution
- Support for multiple questions
- Color-coded responses (Strongly Agree, Agree, Partially Agree)
"""

import matplotlib
# Set backend to 'Agg' for headless environments (Render deployment)
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import numpy as np
import os
import tempfile
from docx.shared import Inches, Cm
from docx import Document
from .document_utils import find_and_replace_text

# Set matplotlib to use a non-interactive backend for server environments
import matplotlib
matplotlib.use('Agg')


def process_feedback_data(request):
    """Extract and process feedback data from form submission."""
    feedback_data = []
    
    # Define the questions (should match frontend)
    questions = [
        "The trainer was able to communicate clearly.",
        "The Content of ECSBC / ENS covered was satisfactory.",
        "Adequate time was provided for question-and-answer session.",
        "The content was appropriately described and key concepts conveyed properly."
    ]
    
    for i, question_text in enumerate(questions, 1):
        question_id = i
        strongly_agree = int(request.form.get(f'question_{question_id}_strongly_agree', '0') or '0')
        agree = int(request.form.get(f'question_{question_id}_agree', '0') or '0')
        partially_agree = int(request.form.get(f'question_{question_id}_partially_agree', '0') or '0')
        
        feedback_data.append({
            'id': question_id,
            'question': question_text,
            'strongly_agree': strongly_agree,
            'agree': agree,
            'partially_agree': partially_agree,
            'total': strongly_agree + agree + partially_agree
        })
    
    return feedback_data


def create_chart_image(question_data, chart_path):
    """Create a horizontal bar chart for a single question."""
    # Set up the figure with professional styling
    plt.style.use('default')
    fig, ax = plt.subplots(figsize=(10, 3), dpi=300)
    
    # Colors matching the frontend design
    colors = ['#10b981', '#3b82f6', '#f59e0b']  # Green, Blue, Orange
    labels = ['Strongly Agree', 'Agree', 'Partially Agree']
    
    values = [
        question_data['strongly_agree'],
        question_data['agree'],
        question_data['partially_agree']
    ]
    
    # Create horizontal bar chart
    y_pos = np.arange(len(labels))
    bars = ax.barh(y_pos, values, color=colors, alpha=0.8, height=0.6)
    
    # Customize the chart
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=11)
    ax.set_xlabel('Number of Responses', fontsize=11, fontweight='bold')
    ax.set_title(f'Q{question_data["id"]}: {question_data["question"]}', 
                fontsize=12, fontweight='bold', pad=20, wrap=True)
    
    # Add value labels on bars
    for i, (bar, value) in enumerate(zip(bars, values)):
        if value > 0:
            ax.text(value + 0.1, bar.get_y() + bar.get_height()/2, 
                   str(value), va='center', fontsize=10, fontweight='bold')
    
    # Styling
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#e5e7eb')
    ax.spines['bottom'].set_color('#e5e7eb')
    ax.grid(axis='x', alpha=0.3, linestyle='--')
    ax.set_axisbelow(True)
    
    # Set x-axis limits with some padding
    max_val = max(values) if values else 1
    ax.set_xlim(0, max_val * 1.2 if max_val > 0 else 5)
    
    # Tight layout and save
    plt.tight_layout()
    plt.savefig(chart_path, dpi=300, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    plt.close()


def generate_feedback_charts(request):
    """Generate all feedback charts and return their file paths."""
    feedback_data = process_feedback_data(request)
    chart_paths = []
    
    # Create temporary directory for charts
    temp_dir = tempfile.mkdtemp()
    
    for question_data in feedback_data:
        if question_data['total'] > 0:  # Only create charts for questions with responses
            chart_filename = f"feedback_chart_q{question_data['id']}.png"
            chart_path = os.path.join(temp_dir, chart_filename)
            
            create_chart_image(question_data, chart_path)
            chart_paths.append(chart_path)
            
            print(f"📊 Generated chart for Question {question_data['id']}: {chart_path}")
    
    return chart_paths


def insert_charts_in_document(doc, chart_paths, placeholder='{{FEEDBACK_CHARTS}}'):
    """Insert generated charts into the Word document at individual placeholders."""
    if not chart_paths:
        print("⚠️ No charts to insert")
        # Remove both old and new placeholder formats
        find_and_replace_text(doc, placeholder, 'No feedback data provided for chart generation.')
        for i in range(1, 5):  # Remove up to 4 individual placeholders
            find_and_replace_text(doc, f'{{{{FEEDBACK_CHART_{i}}}}}', 'No feedback data available.')
        return
    
    # Check if document uses individual chart placeholders
    uses_individual_placeholders = False
    individual_placeholders_found = []
    
    for para in doc.paragraphs:
        para_text = para.text
        for i in range(1, 5):
            placeholder_name = f'FEEDBACK_CHART_{i}'
            if placeholder_name in para_text:
                uses_individual_placeholders = True
                individual_placeholders_found.append(f'{{{{FEEDBACK_CHART_{i}}}}}')
    
    print(f"🔍 Individual placeholders found: {individual_placeholders_found}")
    print(f"📊 Available charts: {len(chart_paths)}")
    
    if uses_individual_placeholders:
        print("✅ Using individual chart placeholder mode")
        # Insert charts at individual placeholders
        for i, chart_path in enumerate(chart_paths):
            if i >= 4:  # Limit to 4 charts max
                print(f"⚠️ Limiting to 4 charts, skipping chart {i+1}")
                break
                
            chart_placeholder = f'{{{{FEEDBACK_CHART_{i+1}}}}}'
            print(f"📍 Processing {chart_placeholder} with chart: {os.path.basename(chart_path)}")
            
            if os.path.exists(chart_path):
                try:
                    # Find the paragraph with this specific placeholder
                    placeholder_found = False
                    for para in doc.paragraphs:
                        if chart_placeholder in para.text:
                            print(f"🎯 Found placeholder {chart_placeholder} in paragraph")
                            
                            # Clear the placeholder text
                            para.clear()
                            
                            # Insert the chart image
                            run = para.add_run()
                            run.add_picture(chart_path, width=Cm(15))  # Full page width
                            
                            # Center the image
                            para.alignment = 1  # Center alignment
                            
                            print(f"✅ Inserted chart {i+1} at {chart_placeholder}: {os.path.basename(chart_path)}")
                            placeholder_found = True
                            break
                    
                    if not placeholder_found:
                        print(f"⚠️ Placeholder {chart_placeholder} not found in document")
                        
                except Exception as e:
                    print(f"❌ Error inserting chart {chart_path}: {str(e)}")
                    # Replace placeholder with error message if it exists
                    find_and_replace_text(doc, chart_placeholder, f"Error loading chart: {os.path.basename(chart_path)}")
            else:
                print(f"❌ Chart file not found: {chart_path}")
                find_and_replace_text(doc, chart_placeholder, "Chart file not found")
        
        # Remove any unused placeholders (charts 3 and 4 if only 2 charts generated)
        for i in range(len(chart_paths) + 1, 5):
            unused_placeholder = f'{{{{FEEDBACK_CHART_{i}}}}}'
            find_and_replace_text(doc, unused_placeholder, '')
            print(f"🧹 Removed unused placeholder: {unused_placeholder}")
            
    else:
        print("✅ Using legacy single placeholder mode")
        # Fall back to old method - single placeholder
        placeholder_found = False
        for para in doc.paragraphs:
            if placeholder in para.text:
                placeholder_found = True
                print(f"🎯 Found legacy placeholder {placeholder}")
                
                # Clear the placeholder text
                para.clear()
                
                # Insert each chart as a new paragraph
                current_para = para
                for i, chart_path in enumerate(chart_paths):
                    if os.path.exists(chart_path):
                        try:
                            # Use the current paragraph for the first chart
                            if i == 0:
                                target_para = current_para
                            else:
                                # Create new paragraph for additional charts
                                target_para = para._parent.add_paragraph()
                            
                            # Insert the chart image
                            run = target_para.add_run()
                            run.add_picture(chart_path, width=Cm(15))  # Full page width
                            
                            # Center the image
                            target_para.alignment = 1  # Center alignment
                            
                            # Add spacing between charts
                            if i < len(chart_paths) - 1:
                                spacing_para = para._parent.add_paragraph()
                                spacing_para.add_run().add_break()
                            
                            print(f"✅ Inserted chart: {os.path.basename(chart_path)}")
                            
                        except Exception as e:
                            print(f"❌ Error inserting chart {chart_path}: {str(e)}")
                            # Add error message to document
                            error_para = para._parent.add_paragraph()
                            error_para.add_run(f"Error loading chart: {os.path.basename(chart_path)}")
                    else:
                        print(f"❌ Chart file not found: {chart_path}")
                
                break
        
        if not placeholder_found:
            print(f"⚠️ Legacy placeholder {placeholder} not found in document")
    
    # Clean up temporary files
    for chart_path in chart_paths:
        try:
            if os.path.exists(chart_path):
                os.remove(chart_path)
        except Exception as e:
            print(f"⚠️ Could not remove temporary chart file {chart_path}: {str(e)}")


def create_summary_feedback_chart(feedback_data, chart_path):
    """Create a summary chart showing all questions together."""
    if not feedback_data or all(q['total'] == 0 for q in feedback_data):
        return None
    
    plt.style.use('default')
    fig, ax = plt.subplots(figsize=(12, 8), dpi=300)
    
    # Prepare data
    questions = [f"Q{q['id']}" for q in feedback_data if q['total'] > 0]
    strongly_agree_data = [q['strongly_agree'] for q in feedback_data if q['total'] > 0]
    agree_data = [q['agree'] for q in feedback_data if q['total'] > 0]
    partially_agree_data = [q['partially_agree'] for q in feedback_data if q['total'] > 0]
    
    x = np.arange(len(questions))
    width = 0.25
    
    # Create grouped bars
    bars1 = ax.bar(x - width, strongly_agree_data, width, label='Strongly Agree', 
                   color='#10b981', alpha=0.8)
    bars2 = ax.bar(x, agree_data, width, label='Agree', 
                   color='#3b82f6', alpha=0.8)
    bars3 = ax.bar(x + width, partially_agree_data, width, label='Partially Agree', 
                   color='#f59e0b', alpha=0.8)
    
    # Customize chart
    ax.set_xlabel('Questions', fontsize=12, fontweight='bold')
    ax.set_ylabel('Number of Responses', fontsize=12, fontweight='bold')
    ax.set_title('Feedback Survey Results Summary', fontsize=14, fontweight='bold', pad=20)
    ax.set_xticks(x)
    ax.set_xticklabels(questions)
    ax.legend(loc='upper right')
    
    # Add value labels on bars
    def add_value_labels(bars):
        for bar in bars:
            height = bar.get_height()
            if height > 0:
                ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                       f'{int(height)}', ha='center', va='bottom', fontsize=9)
    
    add_value_labels(bars1)
    add_value_labels(bars2)
    add_value_labels(bars3)
    
    # Styling
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_axisbelow(True)
    
    plt.tight_layout()
    plt.savefig(chart_path, dpi=300, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close()
    
    return chart_path
